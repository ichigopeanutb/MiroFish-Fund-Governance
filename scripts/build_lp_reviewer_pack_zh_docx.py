from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "lp-facing"
ASSETS = DOCS / "assets"
OUT = DOCS / "MiroFish_Fund_Governance_Nontechnical_Reviewer_Pack_ZH.docx"


BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
FONT = "Arial Unicode MS"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=15 if level == 1 else 12, bold=True, color=BLUE)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(text), size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text), size=10.5)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_metric_table(doc):
    rows = [
        ("Events", "18", "是否能形成可追蹤的基金運作 timeline"),
        ("Capital called", "1,150,000", "capital call 是否清楚可檢查"),
        ("Capital paid", "1,150,000", "synthetic LP 是否按時付款"),
        ("Unfunded commitment", "8,850,000", "未出資承諾是否能用來做後續規劃"),
        ("NAV", "400,000", "portfolio value 是否和現金分配分開呈現"),
        ("Distributions", "1,136,000", "distribution / waterfall 邏輯是否可被檢查"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.55), Inches(1.35), Inches(3.6)]
    headers = ["指標", "Synthetic Result", "Reviewer 要看什麼"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, LIGHT_FILL)
        set_run(cell.paragraphs[0].add_run(text), size=10, bold=True)
    for metric, result, meaning in rows:
        cells = table.add_row().cells
        values = [metric, result, meaning]
        for idx, value in enumerate(values):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_run(p.add_run(value), size=9.5)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"].font.size = Pt(11)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(title.add_run("MiroFish Fund Governance Edition"), size=22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_run(subtitle.add_run("非技術基金經理人試用邀請包 - Synthetic Demo"), size=12, color=MUTED)

    add_body(
        doc,
        "這份文件是給不需要 GitHub、不需要技術背景的基金經理人、前輩或 "
        "LP-facing reviewer 看的。所有內容都使用 synthetic data，不是真實基金資料，"
        "也不是法律、稅務、會計、投資或績效建議。",
    )

    add_heading(doc, "你現在看到的是什麼")
    add_body(
        doc,
        "MiroFish Fund Governance Edition 是一個基金治理模擬工具，用來演練 "
        "capital call、fund terms、waterfall、IC / LPAC 決策、LP 溝通、"
        "證據軌跡和 meeting pack 準備。",
    )

    image_path = ASSETS / "demo-business-page-first-screen.png"
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.25))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(caption.add_run("Synthetic fund-governance simulation 完成後的 demo 畫面。"), size=9, color=MUTED)

    doc.add_page_break()
    add_heading(doc, "完成後報告會長什麼樣子")
    add_body(
        doc,
        "完成後的 synthetic run 會產生一份 operating report，讓 reviewer 看見基金運作發生了什麼、"
        "哪些 fund terms 或 governance issues 被觸發，以及進入 LP、IC 或 LPAC 對話前應該討論什麼。",
    )
    add_metric_table(doc)

    add_heading(doc, "這份 sample 會引出的治理問題")
    add_bullet(doc, "Follow-on reserve 決策應該由 IC approval、LPAC consent，還是 manager discretion 處理？")
    add_bullet(doc, "哪些文件段落或 memo excerpts 應該被綁成 approval 前的 evidence？")
    add_bullet(doc, "Distribution summary 是否已經顯示足夠的 waterfall 邏輯，能支持 LP-facing conversation？")
    add_bullet(doc, "哪些 assumptions 應該在下一次 simulation run 前允許使用者調整？")

    add_heading(doc, "想請你幫忙看的三個問題")
    add_number(doc, "如果你是基金經理人或 LP-facing advisor，這份 output 哪一部分最可能在真實對話裡有用？")
    add_number(doc, "哪一部分還不清楚、太技術、或還不夠可信？")
    add_number(doc, "如果要讓你願意把這個介紹給另一位基金經理人、LP、IC member 或 LPAC reviewer，還需要補什麼？")

    add_heading(doc, "邊界")
    add_body(
        doc,
        "這是一份 synthetic product demonstration。它不包含真實 LP 資料、真實基金文件、投資人身份、"
        "銀行資訊、稅務紀錄、法律意見、會計紀錄或投資建議。",
        after=0,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
