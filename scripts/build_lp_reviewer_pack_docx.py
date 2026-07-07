from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs" / "lp-facing"
ASSETS = DOCS / "assets"
OUT = DOCS / "MiroFish_Fund_Governance_Nontechnical_Reviewer_Pack.docx"


BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, size=15 if level == 1 else 12, bold=True, color=BLUE)
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    set_run(p.add_run(text), size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(text), size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text), size=10.5)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_metric_table(doc):
    rows = [
        ("Events", "18", "Traceable fund-operation timeline"),
        ("Capital called", "1,150,000", "Capital call activity is visible"),
        ("Capital paid", "1,150,000", "Synthetic LP paid on time"),
        ("Unfunded commitment", "8,850,000", "Remaining commitment is visible"),
        ("NAV", "400,000", "Portfolio value shown separately"),
        ("Distributions", "1,136,000", "Distribution logic can be inspected"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    widths = [Inches(1.55), Inches(1.35), Inches(3.6)]
    headers = ["Metric", "Synthetic Result", "Reviewer Meaning"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, LIGHT_FILL)
        set_run(cell.paragraphs[0].add_run(text), size=10, bold=True)
    for metric, result, meaning in rows:
        cells = table.add_row().cells
        values = [metric, result, meaning]
        for idx, value in enumerate(values):
            cells[idx].width = widths[idx]
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_run(p.add_run(value), size=9.5)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(title.add_run("MiroFish Fund Governance Edition"), size=22, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    set_run(
        subtitle.add_run("Nontechnical Reviewer Pack - Synthetic Demo"),
        size=12,
        color=MUTED,
    )

    add_body(
        doc,
        "This pack is for a fund manager, senior advisor, or LP-facing reviewer "
        "who does not need GitHub access. It uses synthetic data only and is not "
        "legal, tax, accounting, investment, or performance advice.",
    )

    add_heading(doc, "What You Are Looking At")
    add_body(
        doc,
        "MiroFish Fund Governance Edition is a simulation layer for rehearsing "
        "fund operations: capital calls, fund terms, waterfall, IC / LPAC "
        "decisions, LP communication, evidence trails, and meeting-pack "
        "preparation.",
    )

    image_path = ASSETS / "demo-business-page-first-screen.png"
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(6.5))
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(
            caption.add_run(
                "Demo screen after a synthetic fund-governance simulation run."
            ),
            size=9,
            color=MUTED,
        )

    add_heading(doc, "Sample Completed Report")
    add_body(
        doc,
        "The completed synthetic run produces an operating report that lets a "
        "reviewer inspect what happened, which fund terms or governance issues "
        "were triggered, and what should be discussed before LP or IC / LPAC "
        "review.",
    )
    add_metric_table(doc)

    doc.add_page_break()
    add_heading(doc, "Governance Questions Raised")
    add_bullet(
        doc,
        "Should the follow-on reserve decision require IC approval, LPAC consent, "
        "or manager discretion?",
    )
    add_bullet(
        doc,
        "Which document excerpts should be bound as evidence before approval?",
    )
    add_bullet(
        doc,
        "Does the distribution summary show enough waterfall logic for an "
        "LP-facing conversation?",
    )
    add_bullet(
        doc,
        "Which assumptions should be editable before the next simulation run?",
    )

    add_heading(doc, "Three Questions For The Reviewer")
    add_number(
        doc,
        "If you were a fund manager or LP-facing advisor, which part of this "
        "output would be most useful in a real conversation?",
    )
    add_number(doc, "Which part is unclear, too technical, or not yet credible?")
    add_number(
        doc,
        "What would need to be added before you would introduce this to another "
        "fund manager, LP, IC member, or LPAC reviewer?",
    )

    add_heading(doc, "Boundary")
    add_body(
        doc,
        "This is a synthetic product demonstration. It does not include real LP "
        "data, real fund documents, investor identifiers, bank information, tax "
        "records, legal advice, accounting records, or investment advice.",
        after=0,
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
